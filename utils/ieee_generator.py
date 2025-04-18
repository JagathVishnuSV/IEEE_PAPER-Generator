import logging
import re
import tempfile
from io import BytesIO

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt, Inches, RGBColor

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def to_roman(num: int) -> str:
    roman_map = [
        (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
        (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
        (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
    ]
    result = ""
    for val, sym in roman_map:
        while num >= val:
            result += sym
            num -= val
    return result


def generate_latex_formula_image(latex_code: str) -> str | None:
    try:
        fig, ax = plt.subplots(figsize=(2, 0.5))
        ax.text(0.5, 0.5, f"${latex_code}$", fontsize=14, ha='center', va='center')
        ax.axis('off')
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            plt.savefig(tmp.name, format='png', bbox_inches='tight', transparent=True)
            plt.close(fig)
            return tmp.name
    except Exception as e:
        logger.error(f"Formula rendering failed: {e}")
        return None


def add_hyperlink(paragraph, display_text: str, url: str):
    """Insert a clickable, blue-underlined hyperlink into `paragraph`."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = display_text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_hyperlinks(paragraph, text: str):
    """
    Scan `text` for markdown [text](url) or bare URLs, emit plain text and
    clickable hyperlinks.
    """
    pattern = re.compile(r'\[([^\]]+)\]\((https?://[^\)]+)\)|(https?://[^\s]+)')
    idx = 0
    for m in pattern.finditer(text):
        start, end = m.span()
        if start > idx:
            paragraph.add_run(text[idx:start])
        if m.group(1) and m.group(2):
            disp, url = m.group(1), m.group(2)
        else:
            disp = url = m.group(3)
        add_hyperlink(paragraph, disp, url)
        idx = end
    if idx < len(text):
        paragraph.add_run(text[idx:])


def extract_and_replace_hyperlinks(text: str, start_idx: int = 1) -> tuple[str, list[str]]:
    """
    Replace every markdown [text](url) in `text` with [n], return modified
    text and list of unique URLs found.
    """
    pattern = r'\[([^\]]+)\]\((https?://[^\)]+)\)'
    citations = []
    ref_map = {}
    out = text

    for match in re.finditer(pattern, text):
        _, url = match.groups()
        if url not in ref_map:
            ref_map[url] = start_idx + len(citations)
            citations.append(url)
        num = ref_map[url]
        out = out.replace(match.group(0), f"[{num}]")

    return out, citations


def insert_footnotes(paragraph, content: str):
    """
    Replace [[footnote:Note text]] with “[∗] Note text” inline.
    """
    pattern = r"\[\[footnote:(.*?)\]\]"
    for note in re.findall(pattern, content):
        content = content.replace(f"[[footnote:{note}]]", f"[*] {note}")
    paragraph.add_run(content)


def set_single_column_layout(section):
    sectPr = section._sectPr
    for c in sectPr.xpath('./w:cols'):
        sectPr.remove(c)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '1')
    sectPr.append(cols)


def set_ieee_column_layout(section):
    sectPr = section._sectPr
    for c in sectPr.xpath('./w:cols'):
        sectPr.remove(c)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    sectPr.append(cols)

def format_Subheading(paragraph):
    run = paragraph.runs[0]
    run.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)

def format_heading(paragraph):
    run = paragraph.runs[0]
    run.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def italic_format_heading(paragraph):
    run = paragraph.runs[0]
    run.italic = True
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)

def generate_ieee_paper(data: dict) -> bytes:
    try:
        doc = Document()

        # ——— Global style ——————————————————————————————
        normal = doc.styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(10)

        sec = doc.sections[0]
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        set_single_column_layout(sec)

        # ——— Title & Authors ——————————————————————————————
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(data['title'].upper())
        r.bold = True
        r.font.size = Pt(16)

        for line in (
            ", ".join(data['authors']),
            "; ".join(data['affiliations']),
            ", ".join(data['emails'])
        ):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_section(0)
        set_ieee_column_layout(doc.sections[-1])

        # ——— Abstract ——————————————————————————————
        italic_format_heading(doc.add_paragraph("Abstract—"))
        abs_text = " ".join(data['abstract']) if isinstance(data['abstract'], list) else data['abstract']
        for para in abs_text.split('\n'):
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.bold = True

        # ——— Keywords ——————————————————————————————
        format_Subheading(doc.add_paragraph("Keywords"))
        kw = doc.add_paragraph(", ".join(data['keywords']))
        kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ——— Main Sections ——————————————————————————————
        fig_ct = tbl_ct = 1
        ref_idx = 1
        extracted = []

        for i, sec_data in enumerate(data['sections'], 1):
            roman = to_roman(i)
            format_heading(doc.add_paragraph(f"{roman}. {sec_data['heading'].upper()}"))

            # Content + citations
            content = sec_data.get("content", "").strip()
            if content:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                mod, urls = extract_and_replace_hyperlinks(content, ref_idx)
                insert_footnotes(p, mod)
                p_runs = p.runs  # text now has "[n]" placeholders; no inline links per IEEE
                ref_idx += len(urls)
                extracted.extend(urls)

            # Images
            for img in sec_data.get("images", []):
                doc.add_picture(img["path"], width=Inches(3))
                cap = doc.add_paragraph(f"Fig. {fig_ct}: {img['caption']}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_ct += 1

            # Tables
            for table in sec_data.get("tables", []):
                tbl = doc.add_table(rows=len(table), cols=len(table[0]))
                tbl.style = "Table Grid"
                for r, row in enumerate(table):
                    for c, val in enumerate(row):
                        tbl.cell(r, c).text = str(val)
                doc.add_paragraph(f"Table {tbl_ct}: Data Table").alignment = WD_ALIGN_PARAGRAPH.CENTER
                tbl_ct += 1

            # Formulas
            for f_i, formula in enumerate(sec_data.get("formulas", []), 1):
                path = generate_latex_formula_image(formula)
                if path:
                    doc.add_picture(path, width=Inches(2))
                    cap = doc.add_paragraph(f"Equation {i}.{f_i}: {formula}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subsections
            for j, sub in enumerate(sec_data.get("subsections", []), 1):
                format_Subheading(doc.add_paragraph(f"{chr(64+j)}. {sub['heading']}"))

                cnt = sub.get("content", "").strip()
                if cnt:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    mod, urls = extract_and_replace_hyperlinks(cnt, ref_idx)
                    insert_footnotes(p, mod)
                    ref_idx += len(urls)
                    extracted.extend(urls)

                for img in sub.get("images", []):
                    doc.add_picture(img["path"], width=Inches(3))
                    cap = doc.add_paragraph(f"Fig. {fig_ct}: {img['caption']}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fig_ct += 1

                for table in sub.get("tables", []):
                    tbl = doc.add_table(rows=len(table), cols=len(table[0]))
                    tbl.style = "Table Grid"
                    for r, row in enumerate(table):
                        for c, val in enumerate(row):
                            tbl.cell(r, c).text = str(val)
                    doc.add_paragraph(f"Table {tbl_ct}: Data Table").alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tbl_ct += 1

                for f_i, formula in enumerate(sub.get("formulas", []), 1):
                    path = generate_latex_formula_image(formula)
                    if path:
                        doc.add_picture(path, width=Inches(2))
                        cap = doc.add_paragraph(f"Equation {roman}.{j}.{f_i}: {formula}")
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ——— References ——————————————————————————————
        # Strip leading [n] from manual refs
        manual = []
        for r in data.get('references', []):
            m = re.match(r'^\[\d+\]\s*(.*)$', r)
            manual.append(m.group(1) if m else r)

        # Dedupe extracted
        uniq = list(dict.fromkeys(extracted))
        filtered = [u for u in uniq if not any(u in mr for mr in manual)]

        combined = manual + [f"[Online]. Available: {u}" for u in filtered]

        format_heading(doc.add_paragraph("REFERENCES"))
        for idx, ref in enumerate(combined, 1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Number + hyperlink any URLs
            add_hyperlinks(p, f"[{idx}] {ref}")

        # ——— Appendix ——————————————————————————————
        if data.get('appendix'):
            format_heading(doc.add_paragraph("Appendix"))
            for item in data['appendix']:
                p = doc.add_paragraph(item)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ——— Write out ——————————————————————————————
        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return out.read()

    except Exception as e:
        logger.error(f"Error generating IEEE paper: {e}")
        raise
